import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from math import pi
from docx import Document
from docx.shared import Inches
import tempfile

# Define questionnaire sections and questions
questionnaire = {
    "Fruits and Vegetables": [
        "How often do you eat at least 5 servings of fruits and vegetables per day?",
        "How often do you include a variety of colors in your fruit and vegetable choices?",
        "How often do you eat fresh fruits instead of fruit juices?",
        "How often do you consume starchy vegetables (e.g., potatoes, corn) compared to non-starchy ones (e.g., spinach, broccoli)?"
    ],
    "Whole Grains": [
        "How often do you choose whole grains (e.g., brown rice, oats, whole wheat bread) instead of refined grains?",
        "How often do you check labels for whole grain ingredients when buying packaged foods?",
        "How often do you eat fiber-rich cereals or grains daily?"
    ],
    "Proteins": [
        "How often do you eat lean protein sources, such as fish, chicken, or legumes?",
        "How often do you limit red meat consumption to once or twice a week?",
        "How often do you consume plant-based protein sources (e.g., beans, lentils, tofu)?",
        "How often do you choose low-fat dairy products or alternatives?"
    ],
    "Fats and Oils": [
        "How often do you limit your intake of saturated fats (e.g., butter, full-fat dairy)?",
        "How often do you use healthy fats, such as olive oil or avocado, in cooking?",
        "How often do you avoid trans fats (e.g., hydrogenated oils in baked goods)?"
    ],
    "Sugar and Sweeteners": [
        "How often do you limit your intake of added sugars in foods and drinks?",
        "How often do you drink water instead of sugary beverages (e.g., soda, energy drinks)?",
        "How often do you check for hidden sugars in packaged or processed foods?"
    ],
    "Sodium and Salt": [
        "How often do you limit your use of table salt during meals?",
        "How often do you choose low-sodium options when buying canned or processed foods?",
        "How often do you flavor your meals with herbs and spices instead of salt?"
    ],
    "Meal Planning and Portion Control": [
        "How often do you plan your meals to include all food groups?",
        "How often do you practice portion control during meals?",
        "How often do you avoid overeating by serving appropriate portions?"
    ],
    "Frequency and Balance": [
        "How often do you eat three balanced meals per day?",
        "How often do you include a healthy snack between meals if needed?",
        "How often do you avoid skipping meals?"
    ],
    "Beverage Choices": [
        "How often do you drink at least 8 glasses of water per day?",
        "How often do you limit your consumption of caffeinated beverages to 2–3 servings per day?",
        "How often do you choose unsweetened beverages, such as herbal teas or plain coffee?"
    ],
    "Food Labels and Awareness": [
        "How often do you read nutrition labels to check for unhealthy ingredients?",
        "How often do you choose foods with low saturated fat, sugar, and sodium content?",
        "How often do you look for foods fortified with essential nutrients (e.g., calcium, vitamin D)?"
    ]
}

# Scoring options
options = ["Always", "Often", "Sometimes", "Rarely", "Never"]

def calculate_section_scores(responses):
    return {section: sum(1 for answer in answers if answer in ["Always", "Often"]) for section, answers in responses.items()}

def visualize_radar_chart(section_scores):
    categories = list(section_scores.keys())
    values = list(section_scores.values())
    values += values[:1]  # Repeat the first value to close the circle

    num_vars = len(categories)
    angles = [n / float(num_vars) * 2 * pi for n in range(num_vars)]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.set_theta_offset(pi / 2)
    ax.set_theta_direction(-1)

    plt.xticks(angles[:-1], categories, color='grey', size=10)
    ax.plot(angles, values, linewidth=2, linestyle='solid', color='blue')
    ax.fill(angles, values, color='blue', alpha=0.25)
    ax.set_title("Dietary Adherence Radar Chart", size=15)
    return fig

def save_to_word(responses, section_scores, radar_chart_path):
    doc = Document()
    doc.add_heading("Dietary Guidelines Adherence Index (DGAI) Report", level=1)
    doc.add_paragraph("Below are your responses and a radar chart visualization based on the DGAI.")

    # Add responses
    doc.add_heading("Your Responses", level=2)
    for section, answers in responses.items():
        doc.add_heading(section, level=3)
        for question, answer in zip(questionnaire[section], answers):
            doc.add_paragraph(f"{question}: {answer}")

    # Add radar chart
    doc.add_heading("Radar Chart", level=2)
    doc.add_picture(radar_chart_path, width=Inches(4.0))

    # Save document to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

def main():
    st.title("Dietary Guidelines Adherence Index (DGAI)")

    st.markdown("""
    **Purpose:**  
    This questionnaire evaluates your adherence to dietary guidelines. Select the option that best describes your behavior for each statement.
    """)

    # User responses
    responses = {}

    for section, questions in questionnaire.items():
        st.header(section)
        responses[section] = []
        for question in questions:
            response = st.selectbox(question, options, key=f"{section}_{question}")
            responses[section].append(response)

    if st.button("Submit Responses"):
        section_scores = calculate_section_scores(responses)

        # Radar Chart Visualization
        fig = visualize_radar_chart(section_scores)
        st.pyplot(fig)

        # Save radar chart as an image
        radar_chart_path = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name
        fig.savefig(radar_chart_path, bbox_inches="tight")

        # Save to Word Document
        word_file = save_to_word(responses, section_scores, radar_chart_path)
        st.success("Word document with radar chart generated successfully!")
        with open(word_file, "rb") as file:
            st.download_button("Download Word Document", file, file_name="DGAI_Report.docx")

if __name__ == "__main__":
    main()
