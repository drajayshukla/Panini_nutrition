import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
from math import pi
from docx import Document
from docx.shared import Inches
import tempfile

# Activity Level Multipliers
activity_levels = {
    "Sedentary (little or no exercise)": 1.2,
    "Lightly Active (light exercise/sports 1-3 days/week)": 1.375,
    "Moderately Active (moderate exercise/sports 3-5 days/week)": 1.55,
    "Very Active (hard exercise/sports 6-7 days/week)": 1.725,
    "Super Active (very hard exercise, physical job, or training)": 1.9
}

# Macronutrient Recommendations
macronutrient_ratios = {
    "Protein": 20,  # Percentage of total calories
    "Fat": 30,
    "Carbohydrates": 50
}


# Calculate BMR
def calculate_bmr(gender, weight, height, age):
    if gender == "Male":
        return 10 * weight + 6.25 * height - 5 * age + 5  # Mifflin-St Jeor Equation for Men
    else:
        return 10 * weight + 6.25 * height - 5 * age - 161  # Mifflin-St Jeor Equation for Women


# Calculate Total Daily Energy Expenditure (TDEE)
def calculate_tdee(bmr, activity_level):
    return bmr * activity_levels[activity_level]


# Calculate Macronutrient Distribution
def calculate_macronutrients(total_calories):
    protein_cal = (macronutrient_ratios["Protein"] / 100) * total_calories
    fat_cal = (macronutrient_ratios["Fat"] / 100) * total_calories
    carb_cal = (macronutrient_ratios["Carbohydrates"] / 100) * total_calories

    return {
        "Protein (g)": protein_cal / 4,  # 1 gram of protein = 4 kcal
        "Fat (g)": fat_cal / 9,  # 1 gram of fat = 9 kcal
        "Carbohydrates (g)": carb_cal / 4  # 1 gram of carbs = 4 kcal
    }


# Generate Word Report
def generate_word_report(inputs, total_calories, macronutrients, images):
    doc = Document()
    doc.add_heading("Daily Calorie Needs and Macronutrient Distribution Report", level=1)

    doc.add_heading("User Details", level=2)
    for key, value in inputs.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("Daily Calorie Needs", level=2)
    doc.add_paragraph(f"Total Calories: {total_calories:.2f} kcal/day")

    doc.add_heading("Macronutrient Distribution", level=2)
    for key, value in macronutrients.items():
        doc.add_paragraph(f"{key}: {value:.2f} g")

    doc.add_heading("Visualizations", level=2)
    for title, image_path in images.items():
        doc.add_heading(title, level=3)
        doc.add_picture(image_path, width=Inches(5.0))

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


# Plot Visualizations
def plot_visualizations(macronutrients):
    # Bar Chart
    categories = list(macronutrients.keys())
    values = list(macronutrients.values())
    fig1, ax1 = plt.subplots(figsize=(10, 5))
    ax1.bar(categories, values, color='skyblue')
    ax1.set_title('Macronutrient Distribution (g)', fontsize=16)
    ax1.set_ylabel('Grams', fontsize=12)
    ax1.set_xticks(range(len(categories)))
    ax1.set_xticklabels(categories, rotation=45, ha='right', fontsize=10)
    plt.tight_layout()
    st.pyplot(fig1)
    bar_chart_path = save_plot_as_image(fig1)

    # Pie Chart
    fig2, ax2 = plt.subplots()
    ax2.pie(values, labels=categories, autopct='%1.1f%%', startangle=140)
    ax2.set_title('Macronutrient Composition (%)', fontsize=14)
    st.pyplot(fig2)
    pie_chart_path = save_plot_as_image(fig2)

    # Radar Chart
    fig3 = radar_chart(macronutrients)
    st.pyplot(fig3)
    radar_chart_path = save_plot_as_image(fig3)

    return bar_chart_path, pie_chart_path, radar_chart_path


def radar_chart(macronutrients):
    labels = list(macronutrients.keys())
    stats = list(macronutrients.values())
    stats += stats[:1]

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig = plt.figure(figsize=(6, 6))
    ax = fig.add_subplot(111, polar=True)
    ax.fill(angles, stats, color='blue', alpha=0.25)
    ax.plot(angles, stats, color='blue', linewidth=2)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_title('Macronutrient Radar Chart', fontsize=14)
    return fig


# Save Plot as Image
def save_plot_as_image(fig):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(temp_file.name, bbox_inches="tight")
    return temp_file.name


# Gamification
def gamification(total_calories):
    if total_calories < 2000:
        return "Bronze", "Consider maintaining a balanced diet to meet your energy needs."
    elif total_calories < 2500:
        return "Silver", "Great! You’re on track to meeting your daily needs."
    else:
        return "Gold", "Excellent! You’re achieving an optimal caloric balance."


# Main Function
def main():
    st.title("Daily Calorie Needs and Macronutrient Distribution Calculator")
    st.markdown("""
    This tool calculates your daily calorie needs and macronutrient distribution based on your goals and activity level.
    """)

    # Inputs
    st.header("User Information")
    gender = st.selectbox("Gender", ["Male", "Female"])
    age = st.number_input("Age (years)", min_value=10, max_value=100, value=30)
    weight = st.number_input("Weight (kg)", min_value=30.0, max_value=200.0, value=70.0)
    height = st.number_input("Height (cm)", min_value=100.0, max_value=250.0, value=170.0)
    activity_level = st.selectbox("Activity Level", list(activity_levels.keys()))
    goal = st.selectbox("Goal", ["Weight Loss", "Maintenance", "Weight Gain"])

    # Calculate BMR and TDEE
    bmr = calculate_bmr(gender, weight, height, age)
    tdee = calculate_tdee(bmr, activity_level)

    # Adjust TDEE Based on Goal
    if goal == "Weight Loss":
        total_calories = tdee - 500
    elif goal == "Weight Gain":
        total_calories = tdee + 500
    else:
        total_calories = tdee

    # Calculate Macronutrient Distribution
    macronutrients = calculate_macronutrients(total_calories)

    # Display Results
    st.header("Results")
    st.subheader("Total Calorie Needs")
    st.write(f"**{total_calories:.2f} kcal/day**")

    st.subheader("Macronutrient Distribution")
    st.json(macronutrients)

    # Gamification
    badge, message = gamification(total_calories)
    st.write(f"**Badge:** {badge}")
    st.success(message)

    # Visualizations
    st.subheader("Visualizations")
    bar_chart_path, pie_chart_path, radar_chart_path = plot_visualizations(macronutrients)

    # Generate Word Report
    if st.button("Download Report"):
        inputs = {
            "Gender": gender,
            "Age": age,
            "Weight": weight,
            "Height": height,
            "Activity Level": activity_level,
            "Goal": goal
        }
        images = {
            "Bar Chart": bar_chart_path,
            "Pie Chart": pie_chart_path,
            "Radar Chart": radar_chart_path
        }
        word_file = generate_word_report(inputs, total_calories, macronutrients, images)
        with open(word_file, "rb") as file:
            st.download_button("Download Report", file, file_name="Calorie_Needs_Report.docx")


if __name__ == "__main__":
    main()
