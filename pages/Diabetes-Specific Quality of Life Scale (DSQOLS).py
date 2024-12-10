import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
from math import pi
from docx import Document
from docx.shared import Inches
import tempfile

# Define questionnaire
questionnaire = {
    "Emotional Impact": [
        "I feel stressed about managing my blood sugar levels.",
        "I feel frustrated about the restrictions diabetes places on my life.",
        "I worry about long-term complications of diabetes.",
        "I feel guilty when my blood sugar levels are not well-controlled.",
        "I feel confident in my ability to manage diabetes (reverse-scored)."
    ],
    "Dietary Behavior": [
        "I feel restricted by the dietary requirements of managing diabetes.",
        "I find it difficult to resist foods that I should avoid.",
        "I feel satisfied with the meal plans I follow for diabetes management.",
        "I struggle to balance my diet with my insulin or medication routine.",
        "I feel confident in making healthy food choices."
    ],
    "Treatment Satisfaction": [
        "I feel satisfied with the effectiveness of my diabetes treatment plan.",
        "I feel overwhelmed by the complexity of diabetes management.",
        "I trust the advice and support I receive from my healthcare provider.",
        "I feel bothered by the frequency of blood sugar checks and insulin injections.",
        "I am satisfied with the tools and technology available to manage my diabetes."
    ],
    "Social and Interpersonal Impact": [
        "I feel diabetes interferes with my social life and activities.",
        "I avoid social gatherings because of diabetes-related dietary restrictions.",
        "I feel supported by my family and friends in managing my diabetes.",
        "I find it difficult to explain my diabetes to others.",
        "I feel judged by others because of my diabetes."
    ],
    "Physical Health": [
        "I feel physically limited because of diabetes.",
        "I experience fatigue that interferes with my daily activities.",
        "I feel energetic and healthy despite having diabetes (reverse-scored).",
        "I have trouble sleeping because of diabetes-related issues.",
        "I feel pain or discomfort related to diabetes (e.g., neuropathy, injections)."
    ],
    "Work and Financial Impact": [
        "I feel that diabetes affects my work performance.",
        "I feel that managing diabetes is financially burdensome.",
        "I have to take time off work due to diabetes-related health issues.",
        "I feel my career choices are limited because of diabetes.",
        "I can manage diabetes effectively without disrupting my work routine (reverse-scored)."
    ]
}

# Scoring system
reverse_scored_questions = [
    "I feel confident in my ability to manage diabetes (reverse-scored).",
    "I feel satisfied with the meal plans I follow for diabetes management.",
    "I feel confident in making healthy food choices.",
    "I feel energetic and healthy despite having diabetes (reverse-scored).",
    "I can manage diabetes effectively without disrupting my work routine (reverse-scored)."
]


def calculate_scores(responses):
    section_scores = {}
    total_score = 0

    for section, questions in responses.items():
        section_score = 0
        for question, score in questions.items():
            if question in reverse_scored_questions:
                section_score += 6 - score  # Reverse scoring
            else:
                section_score += score
        avg_score = section_score / len(questions)
        section_scores[section] = avg_score
        total_score += avg_score

    overall_score = total_score / len(responses)
    return section_scores, overall_score


def plot_visualizations(section_scores):
    # Bar Chart
    sections = list(section_scores.keys())
    scores = list(section_scores.values())
    fig1, ax1 = plt.subplots()
    ax1.bar(sections, scores, color='skyblue')
    ax1.set_title('Section-wise Average Scores')
    ax1.set_ylabel('Average Score')
    st.pyplot(fig1)
    bar_chart_path = save_plot_as_image(fig1)

    # Pie Chart
    fig2, ax2 = plt.subplots()
    ax2.pie(scores, labels=sections, autopct='%1.1f%%', startangle=140)
    ax2.set_title('Score Distribution')
    st.pyplot(fig2)
    pie_chart_path = save_plot_as_image(fig2)

    # Radar Chart
    fig3 = radar_chart(section_scores)
    st.pyplot(fig3)
    radar_chart_path = save_plot_as_image(fig3)

    return bar_chart_path, pie_chart_path, radar_chart_path


def radar_chart(section_scores):
    labels = list(section_scores.keys())
    stats = list(section_scores.values())
    stats += stats[:1]

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig = plt.figure(figsize=(6, 6))
    ax = fig.add_subplot(111, polar=True)
    ax.fill(angles, stats, color='blue', alpha=0.25)
    ax.plot(angles, stats, color='blue', linewidth=2)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels)
    ax.set_title('DSQOLS Radar Chart')
    return fig


def save_plot_as_image(fig):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(temp_file.name, bbox_inches="tight")
    return temp_file.name


def generate_word_report(responses, section_scores, overall_score, badge, message, images):
    doc = Document()
    doc.add_heading('Diabetes-Specific Quality of Life Scale (DSQOLS) Report', level=1)
    doc.add_paragraph(f"Overall Score: {overall_score:.2f}")
    doc.add_paragraph(f"Awarded Badge: {badge}")
    doc.add_paragraph(f"Message: {message}")

    doc.add_heading('Section-wise Scores', level=2)
    for section, score in section_scores.items():
        doc.add_paragraph(f"{section}: {score:.2f}")

    doc.add_heading('Your Responses', level=2)
    for section, questions in responses.items():
        doc.add_heading(section, level=3)
        for question, score in questions.items():
            doc.add_paragraph(f"{question}: {score}")

    doc.add_heading('Visualizations', level=2)
    for title, image_path in images.items():
        doc.add_heading(title, level=3)
        doc.add_picture(image_path, width=Inches(5.0))

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


def gamification(overall_score):
    if overall_score <= 2:
        return "Gold", "Great quality of life! Keep up the excellent work!"
    elif overall_score <= 3.5:
        return "Silver", "Good quality of life. Consider minor improvements for even better results."
    else:
        return "Bronze", "There is room for improvement. Focus on areas that matter most."


def main():
    st.title("Diabetes-Specific Quality of Life Scale (DSQOLS)")
    st.markdown("""
    **Purpose:**  
    This tool assesses how diabetes affects your quality of life.  
    Respond to each question based on the past 4 weeks.
    """)

    # Collect responses
    responses = {}
    for section, questions in questionnaire.items():
        st.header(section)
        responses[section] = {}
        for question in questions:
            score = st.slider(question, min_value=1, max_value=5, value=3, key=f"{section}_{question}")
            responses[section][question] = score

    if st.button("Submit"):
        section_scores, overall_score = calculate_scores(responses)
        badge, message = gamification(overall_score)

        st.subheader("Your Results")
        st.write(f"Overall Score: {overall_score:.2f}")
        st.write(f"Badge: {badge}")
        st.success(message)

        # Visualizations
        st.subheader("Visualizations")
        bar_chart_path, pie_chart_path, radar_chart_path = plot_visualizations(section_scores)

        # Generate Word report
        images = {
            "Bar Chart": bar_chart_path,
            "Pie Chart": pie_chart_path,
            "Radar Chart": radar_chart_path
        }
        word_file = generate_word_report(responses, section_scores, overall_score, badge, message, images)
        st.success("Report generated successfully!")
        with open(word_file, "rb") as file:
            st.download_button("Download Report", file, file_name="DSQOLS_Report.docx")


if __name__ == "__main__":
    main()
