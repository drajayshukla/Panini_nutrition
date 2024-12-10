import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from math import pi
from docx import Document
import tempfile

# Define questionnaire sections and questions
questionnaire = {
    "Awareness": [
        "I notice when I’m full and stop eating.",
        "I pay attention to the flavors and textures of my food while eating.",
        "I eat slowly and savor each bite.",
        "I am aware of how different foods affect my body and energy levels.",
        "I focus solely on my food without distractions like TV or phone."
    ],
    "Emotional Eating": [
        "I eat when I feel stressed.",
        "I eat when I’m bored, even if I’m not hungry.",
        "I eat to cope with negative emotions, such as sadness or anger.",
        "I crave certain foods when I’m upset.",
        "I feel guilty after eating when I wasn’t truly hungry."
    ],
    "Disinhibition": [
        "I find it hard to stop eating even when I’m full.",
        "I eat more when I’m at social gatherings or parties.",
        "I eat past the point of fullness because the food tastes good.",
        "I have difficulty controlling how much I eat of certain foods (e.g., sweets, chips)."
    ],
    "External Cues": [
        "I eat just because food is available, not because I’m hungry.",
        "I eat more when I see others eating.",
        "I eat because it’s 'time to eat,' even if I’m not hungry.",
        "I eat more when I see advertisements or photos of food."
    ],
    "Hunger and Satiety": [
        "I eat only when I’m physically hungry.",
        "I stop eating when I feel satisfied, not overly full.",
        "I can distinguish between physical hunger and emotional hunger."
    ],
    "Eating Environment": [
        "I ensure my meals are eaten in a calm and relaxing environment.",
        "I set aside time specifically for eating without multitasking.",
        "I prefer eating meals at a dining table rather than on the couch or in bed."
    ]
}

# Scoring options
options = ["Always", "Often", "Sometimes", "Rarely", "Never"]

def calculate_scores(responses):
    scores = {
        "High Mindful Eating": 0,
        "Moderate Mindful Eating": 0,
        "Low Mindful Eating": 0
    }

    for section, answers in responses.items():
        for response in answers:
            if response in ["Always", "Often"]:
                scores["High Mindful Eating"] += 1
            elif response == "Sometimes":
                scores["Moderate Mindful Eating"] += 1
            else:
                scores["Low Mindful Eating"] += 1
    return scores

def visualize_bar_chart(scores):
    categories = list(scores.keys())
    values = list(scores.values())

    fig, ax = plt.subplots()
    ax.bar(categories, values, color=["green", "orange", "red"])
    ax.set_title("Mindful Eating Scores")
    ax.set_ylabel("Number of Responses")
    st.pyplot(fig)

def visualize_pie_chart(scores):
    labels = scores.keys()
    values = scores.values()

    fig, ax = plt.subplots()
    ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=140, colors=["green", "orange", "red"])
    ax.set_title("Mindful Eating Response Distribution")
    st.pyplot(fig)

def visualize_radar_chart(responses):
    section_scores = {section: sum(1 for answer in answers if answer in ["Always", "Often"]) for section, answers in responses.items()}
    categories = list(section_scores.keys())
    values = list(section_scores.values())
    values += values[:1]  # Repeat the first value to close the circle

    num_vars = len(categories)
    angles = [n / float(num_vars) * 2 * pi for n in range(num_vars)]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.set_theta_offset(pi / 2)
    ax.set_theta_direction(-1)

    plt.xticks(angles[:-1], categories, color='grey', size=10)
    ax.plot(angles, values, linewidth=2, linestyle='solid', color='blue')
    ax.fill(angles, values, color='blue', alpha=0.25)
    st.pyplot(fig)

def save_to_word(responses, scores, recommendations):
    doc = Document()
    doc.add_heading("Mindful Eating Questionnaire (MEQ) Results", level=1)
    doc.add_paragraph("Below are your responses, scores, and recommendations based on the MEQ.")

    # Add responses
    doc.add_heading("Your Responses", level=2)
    for section, answers in responses.items():
        doc.add_heading(section, level=3)
        for question, answer in zip(questionnaire[section], answers):
            doc.add_paragraph(f"{question}: {answer}")

    # Add scores
    doc.add_heading("Your Scores", level=2)
    for category, score in scores.items():
        doc.add_paragraph(f"{category}: {score}")

    # Add recommendations
    doc.add_heading("Recommendations", level=2)
    for rec in recommendations:
        doc.add_paragraph(rec)

    # Save document to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

def main():
    st.title("Mindful Eating Questionnaire (MEQ)")

    st.markdown("""
    **Purpose:**  
    This questionnaire assesses your awareness, emotional triggers, and mindfulness during meals.  
    Select the option that best describes your behavior for each statement.
    """)

    # User responses
    responses = {}

    for section, questions in questionnaire.items():
        st.header(section)
        responses[section] = []
        for question in questions:
            response = st.selectbox(question, options, key=f"{section}_{question}")
            responses[section].append(response)

    # Calculate scores and display results
    if st.button("Submit Responses"):
        scores = calculate_scores(responses)

        # Display scores
        st.subheader("Your Scores")
        st.write(scores)

        # Visualizations
        st.subheader("Visualizations")
        st.write("### Bar Chart")
        visualize_bar_chart(scores)
        st.write("### Pie Chart")
        visualize_pie_chart(scores)
        st.write("### Radar Chart")
        visualize_radar_chart(responses)

        # Provide recommendations
        recommendations = []
        if scores["High Mindful Eating"] >= scores["Moderate Mindful Eating"] and scores["High Mindful Eating"] >= scores["Low Mindful Eating"]:
            st.success("Great job! You have a high level of mindful eating.")
            recommendations.append("You are doing great with mindful eating. Keep it up!")
        elif scores["Moderate Mindful Eating"] > scores["High Mindful Eating"]:
            st.warning("You have a moderate level of mindful eating. Consider practicing mindful eating techniques.")
            recommendations.append("Consider mindful eating exercises, like eating without distractions or focusing on flavors.")
        else:
            st.error("Your mindful eating level is low. Here are some tips to improve:")
            recommendations.extend([
                "Practice eating without distractions like TV or phones.",
                "Pay attention to your hunger and satiety signals.",
                "Avoid emotional eating triggers by managing stress effectively."
            ])
            for rec in recommendations:
                st.write(f"- {rec}")

        # Save to Word Document
        word_file = save_to_word(responses, scores, recommendations)
        st.success("Word document generated successfully!")
        with open(word_file, "rb") as file:
            st.download_button("Download Results as Word Document", file, file_name="Mindful_Eating_Results.docx")

if __name__ == "__main__":
    main()
