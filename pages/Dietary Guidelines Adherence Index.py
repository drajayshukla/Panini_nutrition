import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from math import pi
from docx import Document
from docx.shared import Inches
import tempfile

# Define the questionnaire
questionnaire = {
    "Fruits and Vegetables": [
        "How often do you eat at least 5 servings of fruits and vegetables per day?",
        "How often do you include a variety of colors in your fruit and vegetable choices?",
        "How often do you eat fresh fruits instead of fruit juices?",
        "How often do you consume starchy vegetables (e.g., potatoes, corn) compared to non-starchy ones (e.g., spinach, broccoli)?"
    ],
    "Whole Grains": [
        "How often do you choose whole grains (e.g., brown rice, oats, whole wheat bread) instead of refined grains?",
        "How often do you check labels for whole grain ingredients when buying packaged foods?",
        "How often do you eat fiber-rich cereals or grains daily?"
    ],
    "Proteins": [
        "How often do you eat lean protein sources, such as fish, chicken, or legumes?",
        "How often do you limit red meat consumption to once or twice a week?",
        "How often do you consume plant-based protein sources (e.g., beans, lentils, tofu)?",
        "How often do you choose low-fat dairy products or alternatives?"
    ],
    "Fats and Oils": [
        "How often do you limit your intake of saturated fats (e.g., butter, full-fat dairy)?",
        "How often do you use healthy fats, such as olive oil or avocado, in cooking?",
        "How often do you avoid trans fats (e.g., hydrogenated oils in baked goods)?"
    ],
    "Sugar and Sweeteners": [
        "How often do you limit your intake of added sugars in foods and drinks?",
        "How often do you drink water instead of sugary beverages (e.g., soda, energy drinks)?",
        "How often do you check for hidden sugars in packaged or processed foods?"
    ],
    "Sodium and Salt": [
        "How often do you limit your use of table salt during meals?",
        "How often do you choose low-sodium options when buying canned or processed foods?",
        "How often do you flavor your meals with herbs and spices instead of salt?"
    ],
    "Meal Planning and Portion Control": [
        "How often do you plan your meals to include all food groups?",
        "How often do you practice portion control during meals?",
        "How often do you avoid overeating by serving appropriate portions?"
    ],
    "Frequency and Balance": [
        "How often do you eat three balanced meals per day?",
        "How often do you include a healthy snack between meals if needed?",
        "How often do you avoid skipping meals?"
    ],
    "Beverage Choices": [
        "How often do you drink at least 8 glasses of water per day?",
        "How often do you limit your consumption of caffeinated beverages to 2–3 servings per day?",
        "How often do you choose unsweetened beverages, such as herbal teas or plain coffee?"
    ],
    "Food Labels and Awareness": [
        "How often do you read nutrition labels to check for unhealthy ingredients?",
        "How often do you choose foods with low saturated fat, sugar, and sodium content?",
        "How often do you look for foods fortified with essential nutrients (e.g., calcium, vitamin D)?"
    ]
}

# Options for answers
options = ["Always", "Often", "Sometimes", "Rarely", "Never"]

# Points for gamification
points_dict = {
    "Always": 5,
    "Often": 4,
    "Sometimes": 3,
    "Rarely": 2,
    "Never": 1
}


def calculate_scores(responses):
    total_score = 0
    section_scores = {}
    for section, answers in responses.items():
        section_score = 0
        for answer in answers:
            section_score += points_dict[answer]
        section_scores[section] = section_score
        total_score += section_score
    return total_score, section_scores


def visualize_scores(section_scores):
    # Bar Chart
    sections = list(section_scores.keys())
    scores = list(section_scores.values())
    fig1, ax1 = plt.subplots(figsize=(10, 5))
    ax1.barh(sections, scores, color='skyblue')
    ax1.set_xlabel('Scores')
    ax1.set_title('Section-wise Scores')
    st.pyplot(fig1)
    # Save the figure
    bar_chart_path = save_plot_as_image(fig1)

    # Pie Chart
    fig2, ax2 = plt.subplots()
    ax2.pie(scores, labels=sections, autopct='%1.1f%%', startangle=140)
    ax2.axis('equal')
    ax2.set_title('Scores Distribution')
    st.pyplot(fig2)
    # Save the figure
    pie_chart_path = save_plot_as_image(fig2)

    # Radar Chart
    fig3 = radar_chart(section_scores)
    st.pyplot(fig3)
    # Save the figure
    radar_chart_path = save_plot_as_image(fig3)

    return bar_chart_path, pie_chart_path, radar_chart_path


def radar_chart(section_scores):
    labels = list(section_scores.keys())
    stats = list(section_scores.values())
    stats += stats[:1]

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig = plt.figure(figsize=(6, 6))
    ax = fig.add_subplot(111, polar=True)
    plt.xticks(angles[:-1], labels)

    ax.plot(angles, stats, color='red', linewidth=2)
    ax.fill(angles, stats, color='red', alpha=0.25)
    ax.set_title('Dietary Guidelines Adherence Radar Chart', y=1.08)

    return fig


def save_plot_as_image(fig):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(temp_file.name, bbox_inches='tight')
    return temp_file.name


def gamification(total_score):
    if total_score >= 180:
        badge = "Gold"
        message = "Excellent adherence to dietary guidelines!"
    elif total_score >= 150:
        badge = "Silver"
        message = "Good job! You're on the right track."
    else:
        badge = "Bronze"
        message = "There's room for improvement. Let's work on it!"
    return badge, message


def generate_word_doc(responses, total_score, section_scores, badge, message, images):
    doc = Document()
    doc.add_heading('Dietary Guidelines Adherence Index (DGAI) Report', 0)
    doc.add_paragraph(f"Total Score: {total_score}")
    doc.add_paragraph(f"Awarded Badge: {badge}")
    doc.add_paragraph(f"Message: {message}")

    doc.add_heading('Section-wise Scores', level=1)
    for section, score in section_scores.items():
        doc.add_paragraph(f"{section}: {score} points")

    doc.add_heading('Visualizations', level=1)
    for title, image_path in images.items():
        doc.add_heading(title, level=2)
        doc.add_picture(image_path, width=Inches(6))

    doc.add_heading('Your Responses', level=1)
    for section, answers in responses.items():
        doc.add_heading(section, level=2)
        for question, answer in zip(questionnaire[section], answers):
            doc.add_paragraph(f"{question}\nYour Answer: {answer}")

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


def main():
    st.title('Dietary Guidelines Adherence Index (DGAI)')
    st.markdown("""
    **Purpose:**  
    The DGAI evaluates how well your diet aligns with established dietary guidelines, focusing on healthy eating patterns and nutrient intake.

    **Instructions:**  
    - Answer each question by selecting the most appropriate option.
    - Use the following scale for most questions:
        - Always
        - Often
        - Sometimes
        - Rarely
        - Never
    """)

    # Collect responses
    responses = {}
    for section, questions in questionnaire.items():
        st.header(section)
        responses[section] = []
        for question in questions:
            answer = st.selectbox(question, options, key=question)
            responses[section].append(answer)

    if st.button('Submit'):
        total_score, section_scores = calculate_scores(responses)
        badge, message = gamification(total_score)

        st.subheader('Your Total Score')
        st.write(f"{total_score} points")
        st.subheader('Badge Awarded')
        st.write(f"{badge} Badge")
        st.success(message)

        # Visualizations
        st.subheader('Visualizations')
        bar_chart_path, pie_chart_path, radar_chart_path = visualize_scores(section_scores)

        # Gamification
        st.balloons()

        # Prepare images for Word doc
        images = {
            "Bar Chart of Section-wise Scores": bar_chart_path,
            "Pie Chart of Scores Distribution": pie_chart_path,
            "Radar Chart of Dietary Adherence": radar_chart_path
        }

        # Generate Word Document
        word_file = generate_word_doc(responses, total_score, section_scores, badge, message, images)
        st.success("Your personalized report has been generated!")
        with open(word_file, "rb") as file:
            st.download_button("Download Report", file, file_name="DGAI_Report.docx")


if __name__ == '__main__':
    main()
