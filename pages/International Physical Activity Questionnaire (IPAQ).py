import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
from math import pi
from docx import Document
from docx.shared import Inches
import tempfile

# Define MET values for activities
MET_VALUES = {
    "vigorous": 8,
    "moderate": 4,
    "walking": 3.3
}


def calculate_activity_score(days, hours, minutes, met):
    total_minutes = (hours * 60) + minutes
    return days * total_minutes * met


def generate_word_report(responses, activity_levels, scores, badge, message, images):
    doc = Document()
    doc.add_heading('International Physical Activity Questionnaire (IPAQ) Report', level=1)

    doc.add_paragraph(f"Activity Level: {activity_levels}")
    doc.add_paragraph(f"Awarded Badge: {badge}")
    doc.add_paragraph(f"Message: {message}")

    doc.add_heading('Your Scores', level=2)
    for activity, score in scores.items():
        doc.add_paragraph(f"{activity.capitalize()} Activity Score: {score} MET-minutes/week")

    doc.add_heading('Your Responses', level=2)
    for section, response in responses.items():
        doc.add_heading(section, level=3)
        for question, answer in response.items():
            doc.add_paragraph(f"{question}: {answer}")

    doc.add_heading('Visualizations', level=2)
    for title, image_path in images.items():
        doc.add_heading(title, level=3)
        doc.add_picture(image_path, width=Inches(5.0))

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


def plot_visualizations(scores):
    # Bar Chart
    activities = list(scores.keys())
    values = list(scores.values())
    fig1, ax1 = plt.subplots()
    ax1.bar(activities, values, color='skyblue')
    ax1.set_title('Physical Activity MET-minutes')
    ax1.set_ylabel('MET-minutes/week')
    st.pyplot(fig1)
    bar_chart_path = save_plot_as_image(fig1)

    # Pie Chart
    fig2, ax2 = plt.subplots()
    ax2.pie(values, labels=activities, autopct='%1.1f%%', startangle=140)
    ax2.set_title('Activity Distribution')
    st.pyplot(fig2)
    pie_chart_path = save_plot_as_image(fig2)

    # Radar Chart
    fig3 = plot_radar_chart(scores)
    st.pyplot(fig3)
    radar_chart_path = save_plot_as_image(fig3)

    return bar_chart_path, pie_chart_path, radar_chart_path


def plot_radar_chart(scores):
    labels = list(scores.keys())
    stats = list(scores.values())
    stats += stats[:1]

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig = plt.figure(figsize=(6, 6))
    ax = fig.add_subplot(111, polar=True)
    ax.fill(angles, stats, color='blue', alpha=0.25)
    ax.plot(angles, stats, color='blue', linewidth=2)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels)
    ax.set_title('Activity Levels Radar Chart')
    return fig


def save_plot_as_image(fig):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(temp_file.name, bbox_inches="tight")
    return temp_file.name


def gamification(total_score):
    if total_score >= 3000:
        badge = "Gold"
        message = "Excellent activity levels! Keep up the great work!"
    elif total_score >= 1500:
        badge = "Silver"
        message = "Good job! Consider increasing your activity slightly for additional health benefits."
    else:
        badge = "Bronze"
        message = "Let's work on improving your activity levels for better health."
    return badge, message


def main():
    st.title("International Physical Activity Questionnaire (IPAQ)")
    st.markdown("""
    **Purpose:**  
    This tool assesses your physical activity levels over the past 7 days.  
    Answer the questions to receive insights into your activity level and recommendations.
    """)

    # Collect responses
    responses = {}
    scores = {}

    st.header("Section 1: Vigorous Physical Activity")
    vigorous_days = st.number_input("Days of vigorous activity in the past week:", min_value=0, max_value=7, value=0)
    vigorous_hours = st.number_input("Hours per day of vigorous activity:", min_value=0, max_value=24, value=0)
    vigorous_minutes = st.number_input("Minutes per day of vigorous activity:", min_value=0, max_value=59, value=0)
    responses["Vigorous Physical Activity"] = {
        "Days": vigorous_days,
        "Hours": vigorous_hours,
        "Minutes": vigorous_minutes
    }
    scores["vigorous"] = calculate_activity_score(vigorous_days, vigorous_hours, vigorous_minutes,
                                                  MET_VALUES["vigorous"])

    st.header("Section 2: Moderate Physical Activity")
    moderate_days = st.number_input("Days of moderate activity in the past week:", min_value=0, max_value=7, value=0)
    moderate_hours = st.number_input("Hours per day of moderate activity:", min_value=0, max_value=24, value=0)
    moderate_minutes = st.number_input("Minutes per day of moderate activity:", min_value=0, max_value=59, value=0)
    responses["Moderate Physical Activity"] = {
        "Days": moderate_days,
        "Hours": moderate_hours,
        "Minutes": moderate_minutes
    }
    scores["moderate"] = calculate_activity_score(moderate_days, moderate_hours, moderate_minutes,
                                                  MET_VALUES["moderate"])

    st.header("Section 3: Walking")
    walking_days = st.number_input("Days of walking in the past week:", min_value=0, max_value=7, value=0)
    walking_hours = st.number_input("Hours per day of walking:", min_value=0, max_value=24, value=0)
    walking_minutes = st.number_input("Minutes per day of walking:", min_value=0, max_value=59, value=0)
    responses["Walking"] = {
        "Days": walking_days,
        "Hours": walking_hours,
        "Minutes": walking_minutes
    }
    scores["walking"] = calculate_activity_score(walking_days, walking_hours, walking_minutes, MET_VALUES["walking"])

    st.header("Section 4: Sitting")
    sitting_hours = st.number_input("Hours per day spent sitting:", min_value=0, max_value=24, value=0)
    sitting_minutes = st.number_input("Minutes per day spent sitting:", min_value=0, max_value=59, value=0)
    responses["Sitting"] = {
        "Hours": sitting_hours,
        "Minutes": sitting_minutes
    }
    scores["sitting"] = (sitting_hours * 60) + sitting_minutes

    # Calculate total activity score
    total_score = scores["vigorous"] + scores["moderate"] + scores["walking"]
    activity_levels = "High" if total_score >= 3000 else "Moderate" if total_score >= 600 else "Low"

    # Display results
    st.subheader("Your Activity Level")
    st.write(f"**{activity_levels} Activity Level**")
    st.write(f"Total Activity Score: {total_score} MET-minutes/week")

    # Gamification
    badge, message = gamification(total_score)
    st.write(f"**Awarded Badge:** {badge}")
    st.success(message)

    # Visualizations
    st.subheader("Visualizations")
    bar_chart_path, pie_chart_path, radar_chart_path = plot_visualizations(scores)

    # Generate Word report
    if st.button("Download Report"):
        images = {
            "Bar Chart": bar_chart_path,
            "Pie Chart": pie_chart_path,
            "Radar Chart": radar_chart_path
        }
        word_file = generate_word_report(responses, activity_levels, scores, badge, message, images)
        with open(word_file, "rb") as file:
            st.download_button("Download Word Report", file, file_name="IPAQ_Report.docx")


if __name__ == "__main__":
    main()
