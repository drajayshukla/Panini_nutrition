import streamlit as st
from docx import Document
import tempfile
import re


def remove_emojis(text):
    emoji_pattern = re.compile(
        "["
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map symbols
        "\U0001F700-\U0001F77F"  # alchemical symbols
        "\U0001F780-\U0001F7FF"  # Geometric Shapes Extended
        "\U0001F800-\U0001F8FF"  # Supplemental Arrows-C
        "\U0001F900-\U0001F9FF"  # Supplemental Symbols and Pictographs
        "\U0001FA00-\U0001FA6F"  # Chess Symbols
        "\U0001FA70-\U0001FAFF"  # Symbols and Pictographs Extended-A
        "\U00002702-\U000027B0"  # Dingbats
        "]+",
        flags=re.UNICODE,
    )
    return emoji_pattern.sub(r"", text)


def save_to_word(data, notes, rewards):
    doc = Document()
    doc.add_heading("Food Diary Summary", level=1)

    # Add food diary data
    for day, meals in data.items():
        doc.add_heading(f"{day}", level=2)
        for meal in meals:
            doc.add_heading(meal["Meal Time"], level=3)
            for key, value in meal.items():
                if key != "Meal Time":
                    doc.add_paragraph(f"{key}: {value}")

    # Add additional notes
    doc.add_heading("Additional Notes", level=2)
    for key, value in notes.items():
        doc.add_paragraph(f"{key}: {value}")

    # Add rewards
    doc.add_heading("Rewards", level=2)
    if rewards:
        for reward in rewards:
            doc.add_paragraph(f"- {reward}")
    else:
        doc.add_paragraph("No rewards earned.")

    # Save the document to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


def main():
    st.title("🍎 Food Diary 🍔")

    st.markdown("""
    **Welcome to Your Food Diary!**  
    Track your meals, earn rewards, and hit streaks for consistent logging!  

    - Earn **Badges** for completing entries.
    - Achieve a **Daily Streak** for finishing all meal logs for the day.
    - Unlock **Bonus Points** by including healthy choices (e.g., fruits, vegetables).
    """)

    # Initialize data storage
    food_diary = {}
    days = ["Day 1"]
    rewards = []

    suggestions = {
        "Breakfast": ["Oatmeal with fruit", "Scrambled eggs and toast", "Smoothie", "Pancakes", "Yogurt with granola"],
        "Mid-morning Snack": ["Coffee", "Fruit", "Protein bar", "Biscuits", "Tea"],
        "Lunch": ["Grilled chicken salad", "Pasta with tomato sauce", "Vegetable stir-fry", "Sandwich",
                  "Rice and lentils"],
        "Afternoon Snack": ["Yogurt", "Nuts", "Fruit juice", "Cookies", "Chips"],
        "Dinner": ["Grilled fish with vegetables", "Soup and bread roll", "Pasta", "Vegetable curry and rice", "Burger"]
    }

    streak_count = 0

    for day in days:
        st.header(f"{day} 🎯")

        completed_meals = 0
        meals = []

        for meal_time in ["Breakfast", "Mid-morning Snack", "Lunch", "Afternoon Snack", "Dinner"]:
            st.subheader(f"{meal_time} 🍽️")
            food_item = st.selectbox(
                f"What did you eat for {meal_time}?",
                suggestions[meal_time] + ["Other"],
                key=f"{day}_{meal_time}_food"
            )
            if food_item == "Other":
                food_item = st.text_input(f"Enter the item for {meal_time}:", key=f"{day}_{meal_time}_food_input")

            portion = st.text_input(
                f"Quantity/Portion for {meal_time} (e.g., '1 bowl', '2 pieces'):",
                key=f"{day}_{meal_time}_portion"
            )
            prep_method = st.selectbox(
                f"Preparation Method for {meal_time}?",
                ["Boiled", "Fried", "Baked", "Grilled", "Raw", "Packaged", "Other"],
                key=f"{day}_{meal_time}_prep"
            )
            location = st.selectbox(
                f"Location for {meal_time}?",
                ["Home", "Office", "Restaurant", "On-the-go", "Other"],
                key=f"{day}_{meal_time}_location"
            )
            notes = st.text_area(f"Notes for {meal_time}?", key=f"{day}_{meal_time}_notes")

            if food_item:
                completed_meals += 1
                meals.append({
                    "Meal Time": meal_time,
                    "Food/Drink Item": food_item,
                    "Quantity/Portion": portion,
                    "Preparation Method": prep_method,
                    "Location": location,
                    "Notes": notes
                })

                # Reward for healthy choices
                if any(word in food_item.lower() for word in ["fruit", "vegetable", "salad", "oatmeal"]):
                    rewards.append(f"Healthy Choice Reward for {meal_time}")

        food_diary[day] = meals

        if completed_meals == 5:
            streak_count += 1
            rewards.append(f"Daily Streak Achieved for {day} 🎉")

    # Progress Tracking
    st.markdown(f"### 🏆 Streak Count: {streak_count}")
    st.progress(streak_count / len(days))

    # Additional Notes Section
    st.header("Additional Notes 📝")
    hunger_levels = st.radio("Were you hungry before meals?", ["Yes", "No"], key="hunger_levels")
    mood = st.selectbox("What was your mood during meals?", ["Stressed", "Relaxed", "Neutral"], key="mood")
    cravings = st.text_input("Did you experience cravings? If yes, for what?", key="cravings")
    hydration = st.slider("How much water did you drink each day? (Approx. in liters)", 0.0, 5.0, 2.0, step=0.1)

    additional_notes = {
        "Hunger Levels": hunger_levels,
        "Mood During Eating": mood,
        "Cravings": cravings,
        "Hydration (Liters)": hydration
    }

    # Display Earned Rewards
    st.header("Your Rewards 🎖️")
    if rewards:
        for reward in rewards:
            st.write(f"- {reward}")
    else:
        st.write("No rewards earned yet! Log more meals to earn badges and rewards.")

    # Generate Word Document
    if st.button("Generate Word Summary"):
        word_file = save_to_word(food_diary, additional_notes, rewards)
        st.success("Word document generated successfully!")
        with open(word_file, "rb") as file:
            st.download_button("Download Word Document", file, file_name="Food_Diary_Summary.docx")


if __name__ == "__main__":
    main()
