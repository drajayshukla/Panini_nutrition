import streamlit as st
import requests
from docx import Document
from docx.shared import Inches
import tempfile
import matplotlib.pyplot as plt
import numpy as np
from math import pi

# Sample Deficiency Questions and Nutrient Mapping
questions = {
    "Do you often feel tired or fatigued?": "Iron",
    "Do you experience frequent muscle cramps or spasms?": "Magnesium",
    "Do you have brittle nails or hair?": "Biotin",
    "Do you experience bone or joint pain?": "Calcium",
    "Do you have dry skin or poor night vision?": "Vitamin A",
    "Do you feel frequent colds or infections?": "Vitamin C",
    "Do you have difficulty concentrating or poor memory?": "Vitamin B12",
}

# Food Suggestions for Nutrients
nutrient_suggestions = {
    "Iron": ["Spinach", "Red meat", "Lentils"],
    "Magnesium": ["Almonds", "Avocado", "Dark chocolate"],
    "Biotin": ["Egg yolk", "Nuts", "Bananas"],
    "Calcium": ["Milk", "Cheese", "Broccoli"],
    "Vitamin A": ["Carrots", "Sweet potatoes", "Pumpkin"],
    "Vitamin C": ["Oranges", "Strawberries", "Bell peppers"],
    "Vitamin B12": ["Fish", "Eggs", "Fortified cereals"],
}

# Generate Word Report
def generate_word_report(user_answers, deficiencies, suggestions, images):
    doc = Document()
    doc.add_heading("Nutrient Deficiency Analysis Report", level=1)

    # User Responses
    doc.add_heading("User Responses", level=2)
    for question, answer in user_answers.items():
        doc.add_paragraph(f"{question}: {answer}")

    # Identified Deficiencies
    doc.add_heading("Identified Deficiencies", level=2)
    for nutrient in deficiencies:
        doc.add_paragraph(f"- {nutrient}")

    # Food Suggestions
    doc.add_heading("Suggested Foods", level=2)
    for nutrient, foods in suggestions.items():
        doc.add_paragraph(f"{nutrient}: {', '.join(foods)}")

    # Visualizations
    doc.add_heading("Visualizations", level=2)
    for title, image_path in images.items():
        doc.add_heading(title, level=3)
        doc.add_picture(image_path, width=Inches(5.0))

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

# Plot Visualizations
def plot_visualizations(deficiencies):
    # Bar Chart
    labels = list(deficiencies.keys())
    values = [1] * len(labels)  # Presence of deficiency
    fig1, ax1 = plt.subplots(figsize=(10, 5))
    ax1.bar(labels, values, color='skyblue')
    ax1.set_title('Identified Deficiencies', fontsize=16)
    ax1.set_ylabel('Presence (1)', fontsize=12)
    ax1.set_xticks(range(len(labels)))
    ax1.set_xticklabels(labels, rotation=45, ha='right', fontsize=10)
    plt.tight_layout()
    bar_chart_path = save_plot_as_image(fig1)

    # Pie Chart
    fig2, ax2 = plt.subplots()
    ax2.pie(values, labels=labels, autopct='%1.1f%%', startangle=140)
    ax2.set_title('Deficiency Distribution', fontsize=14)
    pie_chart_path = save_plot_as_image(fig2)

    # Radar Chart
    fig3 = radar_chart(deficiencies)
    radar_chart_path = save_plot_as_image(fig3)

    return bar_chart_path, pie_chart_path, radar_chart_path

def radar_chart(deficiencies):
    labels = list(deficiencies.keys())
    values = [1] * len(labels)  # Presence of deficiency
    values += values[:1]  # Close the radar chart loop

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color='blue', alpha=0.25)
    ax.plot(angles, values, color='blue', linewidth=2)
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_title("Deficiency Radar Chart", fontsize=14)
    return fig

# Save Plot as Image
def save_plot_as_image(fig):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(temp_file.name, bbox_inches="tight")
    return temp_file.name

# Gamification
def gamification(deficiencies):
    count = len(deficiencies)
    if count == 0:
        return "Gold", "Congratulations! You seem to have no major deficiencies."
    elif count <= 2:
        return "Silver", "Good job! Minor deficiencies detected. Address them with the suggested foods."
    else:
        return "Bronze", "Consider a dietary review to address multiple deficiencies."

# Main App
def main():
    st.title("Nutrient Deficiency Analysis")
    st.markdown("Answer questions to identify potential deficiencies and get food suggestions!")

    # Questionnaire
    user_answers = {}
    st.header("Questionnaire")
    for question, nutrient in questions.items():
        user_answers[question] = st.radio(question, ["Yes", "No"], key=question)

    # Analyze Responses
    deficiencies = {
        nutrient: questions[question]
        for question, answer in user_answers.items()
        if answer == "Yes"
    }

    suggestions = {
        nutrient: nutrient_suggestions[nutrient]
        for nutrient in deficiencies.values()
    }

    # Display Results
    if st.button("Analyze"):
        st.header("Results")
        if deficiencies:
            st.success("Deficiencies Detected!")
            st.write("Identified Deficiencies:")
            st.json(list(deficiencies.values()))
            st.write("Food Suggestions:")
            st.json(suggestions)

            # Gamification
            badge, message = gamification(deficiencies)
            st.subheader("Gamification")
            st.write(f"**Badge Earned:** {badge}")
            st.success(message)

            # Visualizations
            st.header("Visualizations")
            bar_chart_path, pie_chart_path, radar_chart_path = plot_visualizations(deficiencies)

            # Generate Report
            if st.button("Generate Report"):
                images = {
                    "Bar Chart": bar_chart_path,
                    "Pie Chart": pie_chart_path,
                    "Radar Chart": radar_chart_path
                }
                report = generate_word_report(user_answers, deficiencies, suggestions, images)
                with open(report, "rb") as file:
                    st.download_button("Download Report", file, file_name="Deficiency_Report.docx")
        else:
            st.success("No major deficiencies detected! Keep up your healthy diet.")

if __name__ == "__main__":
    main()
