import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
from math import pi
from docx import Document
from docx.shared import Inches
import tempfile

# Constants
activity_levels = {
    "Sedentary": 1.2,
    "Lightly Active": 1.375,
    "Moderately Active": 1.55,
    "Very Active": 1.725,
    "Extra Active": 1.9
}

icmr_rda = {
    "Male": {"Energy": 2425, "Protein (g)": 60, "Fat (g)": 25, "Carbohydrates (g)": 364, "Iron (mg)": 17,
             "Calcium (mg)": 600, "Sodium (mg)": 2000, "Potassium (mg)": 3500},
    "Female": {"Energy": 1875, "Protein (g)": 50, "Fat (g)": 20, "Carbohydrates (g)": 281, "Iron (mg)": 21,
               "Calcium (mg)": 600, "Sodium (mg)": 2000, "Potassium (mg)": 3500},
    "Pregnant": {"Energy": 350 + 1875, "Protein (g)": 65, "Fat (g)": 25, "Carbohydrates (g)": 300, "Iron (mg)": 35,
                 "Calcium (mg)": 1200, "Sodium (mg)": 2000, "Potassium (mg)": 3500},
    "Lactating": {"Energy": 600 + 1875, "Protein (g)": 75, "Fat (g)": 30, "Carbohydrates (g)": 325, "Iron (mg)": 21,
                  "Calcium (mg)": 1200, "Sodium (mg)": 2000, "Potassium (mg)": 3500}
}


def calculate_bmr(age, gender, weight, height):
    if gender == "Male":
        return 10 * weight + 6.25 * height - 5 * age + 5  # Mifflin-St Jeor
    else:
        return 10 * weight + 6.25 * height - 5 * age - 161  # Mifflin-St Jeor


def calculate_energy_expenditure(bmr, activity_level):
    return bmr * activity_levels[activity_level]


def macronutrient_distribution(total_energy, protein_g, fat_g):
    protein_cal = protein_g * 4
    fat_cal = fat_g * 9
    carb_cal = total_energy - (protein_cal + fat_cal)
    return {
        "Protein": protein_cal,
        "Fat": fat_cal,
        "Carbohydrates": carb_cal
    }


def generate_word_report(inputs, rda, energy_expenditure, macronutrients, images):
    doc = Document()
    doc.add_heading('Personalized Diet Recommendation', level=1)
    doc.add_heading('User Details', level=2)
    for key, value in inputs.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading('Energy Expenditure and RDA', level=2)
    for key, value in rda.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.add_paragraph(f"Energy Expenditure: {energy_expenditure:.2f} kcal/day")

    doc.add_heading('Macronutrient Distribution', level=2)
    for key, value in macronutrients.items():
        doc.add_paragraph(f"{key}: {value:.2f} kcal")

    doc.add_heading('Visualizations', level=2)
    for title, image_path in images.items():
        doc.add_heading(title, level=3)
        doc.add_picture(image_path, width=Inches(5.0))

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


def plot_visualizations(rda, macronutrients):
    # Bar Chart
    categories = list(rda.keys())
    values = list(rda.values())
    fig1, ax1 = plt.subplots(figsize=(10, 5))
    ax1.bar(categories, values, color='skyblue')
    ax1.set_title('RDA Breakdown', fontsize=14)
    ax1.set_ylabel('Amount', fontsize=12)
    ax1.set_xticklabels(categories, rotation=45, ha='right', fontsize=10)
    st.pyplot(fig1)
    bar_chart_path = save_plot_as_image(fig1)

    # Pie Chart
    macronutrient_labels = macronutrients.keys()
    macronutrient_values = macronutrients.values()
    fig2, ax2 = plt.subplots()
    ax2.pie(macronutrient_values, labels=macronutrient_labels, autopct='%1.1f%%', startangle=140)
    ax2.set_title('Macronutrient Composition', fontsize=14)
    st.pyplot(fig2)
    pie_chart_path = save_plot_as_image(fig2)

    # Radar Chart
    fig3 = radar_chart(rda)
    st.pyplot(fig3)
    radar_chart_path = save_plot_as_image(fig3)

    return bar_chart_path, pie_chart_path, radar_chart_path


def radar_chart(rda):
    labels = list(rda.keys())
    stats = list(rda.values())
    stats += stats[:1]

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig = plt.figure(figsize=(6, 6))
    ax = fig.add_subplot(111, polar=True)
    ax.fill(angles, stats, color='blue', alpha=0.25)
    ax.plot(angles, stats, color='blue', linewidth=2)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_title('RDA Radar Chart', fontsize=14)
    return fig


def save_plot_as_image(fig):
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(temp_file.name, bbox_inches="tight")
    return temp_file.name


def main():
    st.title("Personalized Diet Recommendation")
    st.markdown("""
    **Purpose:**  
    This tool provides personalized dietary recommendations based on user inputs. It calculates energy expenditure using established formulas and suggests RDAs based on ICMR guidelines.
    """)

    # Inputs
    st.header("User Information")
    age = st.number_input("Age (years)", min_value=1, max_value=100, value=30)
    gender = st.selectbox("Gender", ["Male", "Female"])
    weight = st.number_input("Weight (kg)", min_value=1.0, max_value=200.0, value=70.0)
    height = st.number_input("Height (cm)", min_value=50.0, max_value=250.0, value=170.0)
    activity_level = st.selectbox("Activity Level", list(activity_levels.keys()))
    is_pregnant = st.checkbox("Pregnant") if gender == "Female" else False
    is_lactating = st.checkbox("Lactating") if gender == "Female" else False

    # BMR and Energy Expenditure
    bmr = calculate_bmr(age, gender, weight, height)
    energy_expenditure = calculate_energy_expenditure(bmr, activity_level)

    # RDA Selection
    if is_pregnant:
        rda = icmr_rda["Pregnant"]
    elif is_lactating:
        rda = icmr_rda["Lactating"]
    else:
        rda = icmr_rda[gender]

    # Macronutrient Distribution
    macronutrients = macronutrient_distribution(rda["Energy"], rda["Protein (g)"], rda["Fat (g)"])

    # Display Results
    st.header("Results")
    st.subheader("Energy Expenditure")
    st.write(f"Your estimated energy expenditure: {energy_expenditure:.2f} kcal/day")

    st.subheader("Recommended Dietary Allowances (ICMR)")
    st.write(rda)

    st.subheader("Macronutrient Distribution(Kcal)")
    st.write(macronutrients)

    # Visualizations
    st.subheader("Visualizations")
    bar_chart_path, pie_chart_path, radar_chart_path = plot_visualizations(rda, macronutrients)

    # Generate Word Report
    if st.button("Download Report"):
        inputs = {
            "Age": age,
            "Gender": gender,
            "Weight": weight,
            "Height": height,
            "Activity Level": activity_level,
            "Pregnant": is_pregnant,
            "Lactating": is_lactating
        }
        images = {
            "Bar Chart": bar_chart_path,
            "Pie Chart": pie_chart_path,
            "Radar Chart": radar_chart_path
        }
        word_file = generate_word_report(inputs, rda, energy_expenditure, macronutrients, images)
        with open(word_file, "rb") as file:
            st.download_button("Download Report", file, file_name="Personalized_Diet_Report.docx")


if __name__ == "__main__":
    main()
