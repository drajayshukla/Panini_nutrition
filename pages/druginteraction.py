import streamlit as st
import requests
from docx import Document
import tempfile
import matplotlib.pyplot as plt
import networkx as nx

# RxNorm and OpenFDA API URLs
RXNORM_BASE_URL = "https://rxnav.nlm.nih.gov/REST"
OPENFDA_URL = "https://api.fda.gov/drug/label.json"


# Function to fetch RxCUI (RxNorm Concept Unique Identifier) for a drug
def get_rxcui(drug_name):
    url = f"{RXNORM_BASE_URL}/rxcui.json"
    response = requests.get(url, params={"name": drug_name})
    if response.status_code == 200:
        data = response.json()
        return data.get("idGroup", {}).get("rxnormId", [None])[0]
    return None


# Function to check interactions between drugs using RxNorm
def check_interactions(rxcuis):
    url = f"{RXNORM_BASE_URL}/interaction/list.json"
    response = requests.get(url, params={"rxcuis": "+".join(rxcuis)})
    if response.status_code == 200:
        data = response.json()
        interactions = data.get("fullInteractionTypeGroup", [])
        return interactions
    return None


# Fallback: Fetch drug warnings using OpenFDA
def fetch_drug_warnings(drug_name):
    response = requests.get(OPENFDA_URL, params={"search": f"active_ingredient:{drug_name}", "limit": 1})
    if response.status_code == 200:
        data = response.json()
        results = data.get("results", [])
        if results:
            return results[0].get("warnings", ["No warnings found"])
    return None


# Plot Interaction Network
def plot_interaction_network(interactions):
    graph = nx.Graph()
    for interaction in interactions:
        drugs = interaction.get("minConcept", [])
        for i, drug1 in enumerate(drugs):
            for j, drug2 in enumerate(drugs):
                if i < j:
                    graph.add_edge(drug1["name"], drug2["name"])

    # Plot the graph
    plt.figure(figsize=(10, 10))
    pos = nx.spring_layout(graph)
    nx.draw_networkx_nodes(graph, pos, node_size=500, node_color='lightblue')
    nx.draw_networkx_edges(graph, pos, width=1.5)
    nx.draw_networkx_labels(graph, pos, font_size=10)
    plt.title("Drug Interaction Network", fontsize=16)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    plt.savefig(temp_file.name, bbox_inches='tight')
    return temp_file.name


# Generate Word Report
def generate_word_report(drug_names, interactions, warnings, interaction_graph_path):
    doc = Document()
    doc.add_heading("Drug Interaction Report", level=1)

    # Drug Names
    doc.add_heading("Drugs Checked", level=2)
    for drug in drug_names:
        doc.add_paragraph(f"- {drug}")

    # Interaction Details
    doc.add_heading("Interactions", level=2)
    if interactions:
        for interaction in interactions:
            doc.add_paragraph(interaction)
    else:
        doc.add_paragraph("No significant interactions found.")

    # Warnings
    doc.add_heading("Drug Warnings (Fallback)", level=2)
    for drug, warning in warnings.items():
        doc.add_paragraph(f"{drug}: {warning}")

    # Add Interaction Graph
    if interaction_graph_path:
        doc.add_heading("Interaction Network Graph", level=2)
        doc.add_picture(interaction_graph_path)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


# Main App
def main():
    st.title("Drug Interaction Checker")
    st.markdown("Enter drug names to check for potential interactions.")

    # User Input
    drug_names = st.text_area("Enter drug names (one per line):").splitlines()

    if st.button("Check Interactions"):
        if not drug_names:
            st.error("Please enter at least two drugs to check interactions.")
            return

        # Fetch RxCUIs
        rxcuis = []
        for drug in drug_names:
            rxcui = get_rxcui(drug)
            if rxcui:
                rxcuis.append(rxcui)
            else:
                st.error(f"Unable to find RxCUI for drug: {drug}")

        # Check Interactions
        interactions_data = check_interactions(rxcuis)
        interactions = []
        if interactions_data:
            for group in interactions_data:
                for interaction in group.get("fullInteractionType", []):
                    description = interaction.get("interactionPair", [{}])[0].get("description", "No description")
                    interactions.append(description)

        # Fallback to OpenFDA if no interactions found
        warnings = {}
        if not interactions:
            st.warning("No interactions found. Checking for drug warnings...")
            for drug in drug_names:
                warning = fetch_drug_warnings(drug)
                warnings[drug] = warning

        # Display Results
        if interactions:
            st.subheader("Interaction Results")
            st.write(interactions)

            # Plot Network Graph
            interaction_graph_path = plot_interaction_network(interactions)
            st.image(interaction_graph_path, caption="Interaction Network Graph", use_column_width=True)
        else:
            st.subheader("Drug Warnings (Fallback)")
            st.json(warnings)

        # Generate Report
        if st.button("Download Report"):
            interaction_graph_path = None
            if interactions:
                interaction_graph_path = plot_interaction_network(interactions)
            report = generate_word_report(drug_names, interactions, warnings, interaction_graph_path)
            with open(report, "rb") as file:
                st.download_button("Download Report", file, file_name="Drug_Interaction_Report.docx")


if __name__ == "__main__":
    main()
